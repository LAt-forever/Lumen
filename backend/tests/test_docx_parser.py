from pathlib import Path

import pytest
from docx import Document as DocxDocument

from service.core.parsers.docx_parser import DocxParser
from service.core.parsers.base import ParseResult


class FakeSource:
    def __init__(self, content=None, url=None, filename=None, title="", source_type=""):
        self.content = content
        self.url = url
        self.filename = filename
        self.title = title
        self.source_type = source_type


@pytest.fixture
def parser():
    return DocxParser()


def test_supported_types(parser):
    assert parser.supported_types == frozenset({"docx"})


@pytest.mark.asyncio
async def test_parse_missing_filename_raises(parser):
    source = FakeSource(filename=None)
    with pytest.raises(ValueError, match="missing filename"):
        await parser.parse(source)


@pytest.mark.asyncio
async def test_parse_docx(parser, tmp_path: Path, monkeypatch):
    from service.core import storage
    monkeypatch.setattr(storage, "UPLOAD_ROOT", tmp_path)

    docx_path = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("Document Title")
    doc.add_paragraph("First paragraph of content.")
    doc.add_paragraph("Second paragraph of content.")
    doc.save(str(docx_path))

    source = FakeSource(filename="test.docx")
    result = await parser.parse(source)
    assert isinstance(result, ParseResult)
    assert result.title == "Document Title"
    assert "First paragraph of content." in result.text
    assert "Second paragraph of content." in result.text


@pytest.mark.asyncio
async def test_parse_docx_empty_paragraphs_skipped(parser, tmp_path: Path, monkeypatch):
    from service.core import storage
    monkeypatch.setattr(storage, "UPLOAD_ROOT", tmp_path)

    docx_path = tmp_path / "empty.docx"
    doc = DocxDocument()
    doc.add_paragraph("")
    doc.add_paragraph("Only real text.")
    doc.add_paragraph("")
    doc.save(str(docx_path))

    source = FakeSource(filename="empty.docx")
    result = await parser.parse(source)
    assert result.title == "Only real text."
    assert result.text == "Only real text."
